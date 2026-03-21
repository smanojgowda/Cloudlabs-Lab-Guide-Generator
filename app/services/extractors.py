"""
Extractors for PDF, DOCX, and image files.
"""

import base64
import io

import pdfplumber
from docx import Document
from PIL import Image


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num} ---\n{text}")
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name
        if style.startswith("Heading 1"):
            parts.append(f"# {para.text}")
        elif style.startswith("Heading 2"):
            parts.append(f"## {para.text}")
        elif style.startswith("Heading 3"):
            parts.append(f"### {para.text}")
        elif style.startswith("List"):
            parts.append(f"- {para.text}")
        else:
            parts.append(para.text)
    for table in doc.tables:
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
        if rows:
            parts.append("\n".join(rows))
    return "\n\n".join(parts)


def prepare_image_for_api(file_bytes: bytes, filename: str) -> dict:
    """Resize and base64-encode an image for the vision API."""
    ext = filename.lower().split(".")[-1]
    media_type_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "gif": "image/gif", "webp": "image/webp",
    }
    media_type = media_type_map.get(ext, "image/jpeg")

    img = Image.open(io.BytesIO(file_bytes))
    max_dim = 1568
    if max(img.size) > max_dim:
        ratio = max_dim / max(img.size)
        img = img.resize((int(img.width * ratio), int(img.height * ratio)), Image.LANCZOS)

    output = io.BytesIO()
    fmt = "JPEG" if "jpeg" in media_type else ext.upper()
    if fmt == "JPG":
        fmt = "JPEG"
    img.save(output, format=fmt)
    output.seek(0)

    return {
        "base64": base64.standard_b64encode(output.read()).decode("utf-8"),
        "media_type": media_type,
    }


def read_markdown_file(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")
